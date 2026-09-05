"""The Word pipeline: the theme compiled to styles, and the walk that uses them.

There is no equivalent of ``pypdf`` here, and none is needed: a DOCX file is a
zip of XML, and ``python-docx`` reads back everything that was written. So the
tests open the result and assert on structure — which style a paragraph
carries, what a table's header row does, where a hyperlink points.

Two kinds of assertion sit side by side below. Anything with a public
``python-docx`` API is checked through it. The helpers in ``amethyst.ooxml`` have
no such API by definition, so those are checked against the XML they generate
— which is the whole reason they are small functions in a module of their own.

What none of this proves is that Microsoft Word is happy with the result.
Reading a file back with the library that wrote it shows it is well-formed,
not that Word will accept it; that check has to be made by opening one.
"""

from __future__ import annotations

import io
import re
import subprocess
import sys
from datetime import date
from pathlib import Path

import pytest
from docx import Document as read_docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Emu, Pt

from amethyst.document import Document, load_document
from amethyst.ooxml import (
    BOOKMARK_MAX_LENGTH,
    bookmark_name,
    field,
    link,
    numbering_instance,
    properties_child,
    repeat_as_header,
    set_borders,
    shade,
)
from amethyst.render import RenderOptions, render_docx
from amethyst.render.furniture import CONTENTS_HEADING
from amethyst.theme import default_theme
from amethyst.theme.to_docx import (
    BODY_STYLE,
    CODE_INLINE_STYLE,
    CODE_STYLE,
    COVER_DATE_STYLE,
    COVER_SPACE_AFTER,
    FOOTNOTE_STYLE,
    QUOTE_STYLE,
    SUBTITLE_STYLE,
    TABLE_TEXT_STYLE,
    TITLE_PAGE_OFFSET,
    TITLE_STYLE,
    TOC_HEADING_STYLE,
    TOC_INDENT,
    apply_theme,
    family,
    length,
    page_margins,
    page_size,
)

A4_MM = (210.0, 297.0)
LETTER_MM = (215.9, 279.4)


def convert(text: str, base_dir: Path | None = None, **options: object):
    """Render Markdown to a Word file and open it again."""
    document = Document.from_markdown(text, base_dir=base_dir or Path.cwd())
    return reopen(render_docx(document, RenderOptions(**options)).data)


def reopen(data: bytes):
    return read_docx(io.BytesIO(data))


def styles(document) -> list[str]:
    return [paragraph.style.name for paragraph in document.paragraphs]


def texts(document) -> list[str]:
    return [paragraph.text for paragraph in document.paragraphs]


def warnings_from(text: str, **options: object) -> list[str]:
    collected: list[str] = []
    Document.from_markdown(text)
    render_docx(
        Document.from_markdown(text), RenderOptions(warn=collected.append, **options)
    )
    return collected


def xml_of(element) -> str:
    """An element's XML without the namespace declarations that bury it."""
    return re.sub(r"\sxmlns:[a-z0-9]+=\"[^\"]*\"", "", element.xml)


def blank_docx():
    """An empty Word document with the default theme already applied."""
    document = read_docx()
    apply_theme(document, default_theme())
    return document


def test_the_theme_compiler_can_be_imported_before_anything_else():
    """It could not be, for a whole phase, and nothing noticed.

    ``theme/to_docx.py`` needs the OOXML helpers, and while those lived inside
    ``render/`` importing them ran ``render/__init__``, which imported the
    walker, which imported the theme compiler again — half-built. Every module
    here happens to import ``amethyst.render`` first, which enters the ring
    where it closes, so the suite could not see it. A subprocess can.
    """
    for module in (
        "amethyst.theme.to_docx",
        # Phase 7's additions, which reach across the same layers: the config
        # layer reads the theme layer and the render layer, and the fetcher
        # reads the parse layer and the render layer.
        "amethyst.config",
        "amethyst.remote",
        "amethyst.render.highlight",
    ):
        result = subprocess.run(
            [sys.executable, "-c", f"import {module}"],
            capture_output=True,
            text=True,
        )
        assert result.returncode == 0, f"{module}: {result.stderr}"


# --- the theme, compiled to styles ----------------------------------------


def test_a_font_stack_becomes_the_first_real_family():
    assert family(["Iowan Old Style", "Georgia", "serif"]) == "Iowan Old Style"


def test_generic_families_are_skipped_rather_than_named():
    """Word would read "ui-monospace" as a font nobody has installed."""
    assert family(["ui-monospace", "SF Mono", "monospace"]) == "SF Mono"


def test_an_entirely_generic_stack_still_yields_a_font_word_has():
    assert family(["monospace"]) == "Courier New"


def test_the_body_style_carries_the_themes_type():
    theme = default_theme()
    style = blank_docx().styles[BODY_STYLE]
    assert style.font.name == "Iowan Old Style"
    assert style.font.size == Pt(theme.type.size)
    assert style.paragraph_format.line_spacing == theme.type.line_height


def test_headings_are_unbound_from_words_own_theme_fonts():
    """The bug this prevents is invisible: the font is set and ignored.

    A built-in heading names its font twice, plainly and as a reference to the
    document theme's major font, and the reference is the one Word obeys.
    """
    heading = blank_docx().styles["Heading 1"].element
    fonts = heading.find(f"{qn('w:rPr')}/{qn('w:rFonts')}")
    assert fonts.get(qn("w:ascii")) == "Iowan Old Style"
    assert fonts.get(qn("w:asciiTheme")) is None


def test_headings_take_their_size_from_the_scale():
    """To the nearest half point, which is the smallest size Word can hold."""
    theme = default_theme()
    document = blank_docx()
    for level, scale in enumerate(theme.type.headings, start=1):
        style = document.styles[f"Heading {level}"]
        wanted = Pt(theme.type.size * scale)
        assert abs(int(style.font.size) - int(wanted)) <= int(Pt(0.25)), level


def test_the_last_heading_level_is_muted_and_not_italic():
    """Word's Heading 6 arrives italic; the stylesheet's h6 is muted instead."""
    style = blank_docx().styles["Heading 6"]
    assert str(style.font.color.rgb) == default_theme().colors.muted.lstrip("#").upper()
    assert style.font.italic is False


def test_the_code_style_is_shaded_and_bordered():
    properties = xml_of(blank_docx().styles[CODE_STYLE].element)
    assert 'w:fill="F6F4F8"' in properties
    assert "w:pBdr" in properties


def test_a_custom_style_is_based_on_the_body_one():
    document = blank_docx()
    assert document.styles[FOOTNOTE_STYLE].base_style.name == BODY_STYLE


def test_the_cover_styles_are_unbound_from_the_look_word_ships_them_with():
    """``Title`` arrives ruled off in an accent colour that belongs to no theme
    here, and ``Subtitle`` italic and carrying a stray numbering reference."""
    document = blank_docx()
    theme = default_theme()

    title = document.styles[TITLE_STYLE]
    assert "w:pBdr" not in xml_of(title.element)
    assert "asciiTheme" not in xml_of(title.element)
    assert title.font.size == Pt(round(theme.type.size * theme.type.title * 2) / 2)

    subtitle = document.styles[SUBTITLE_STYLE]
    assert "w:numPr" not in xml_of(subtitle.element)
    assert subtitle.font.italic is False


def test_a_contents_style_exists_for_every_heading_level():
    document = blank_docx()
    for level in range(1, 7):
        assert document.styles[f"TOC {level}"].type == WD_STYLE_TYPE.PARAGRAPH


def test_a_contents_entry_is_ruled_out_to_the_page_number_with_dots():
    """The stylesheet's ``leader('.')`` by another name."""
    document = blank_docx()
    section = document.sections[0]
    stops = document.styles["TOC 1"].paragraph_format.tab_stops
    assert [(stop.position, stop.leader) for stop in stops] == [
        (
            section.page_width - section.left_margin - section.right_margin,
            WD_TAB_LEADER.DOTS,
        )
    ]


def test_contents_entries_are_stepped_in_one_level_at_a_time():
    document = blank_docx()
    theme = default_theme()
    step = Pt(theme.type.size * TOC_INDENT)
    indents = [
        document.styles[f"TOC {level}"].paragraph_format.left_indent
        for level in (1, 2, 3)
    ]
    assert indents == [Pt(0), step, Pt(theme.type.size * TOC_INDENT * 2)]


# --- page geometry --------------------------------------------------------


@pytest.mark.parametrize(
    ("written", "millimetres"),
    [
        ("A4", A4_MM),
        ("letter", LETTER_MM),
        ("A4 landscape", (A4_MM[1], A4_MM[0])),
        ("210mm 297mm", A4_MM),
        ("8.5in 11in", LETTER_MM),
    ],
)
def test_a_page_size_is_read_as_a_width_and_a_height(written, millimetres):
    width, height = page_size(written)
    assert round(Emu(width).mm, 1) == round(millimetres[0], 1)
    assert round(Emu(height).mm, 1) == round(millimetres[1], 1)


@pytest.mark.parametrize("written", ["", "A9", "wide", "10 20", "1cm 2cm 3cm"])
def test_a_page_size_word_cannot_read_is_reported_as_such(written):
    assert page_size(written) is None


@pytest.mark.parametrize(
    ("written", "centimetres"),
    [
        ("2cm", (2, 2, 2, 2)),
        ("2cm 3cm", (2, 3, 2, 3)),
        ("1cm 2cm 3cm", (1, 2, 3, 2)),
        ("1cm 2cm 3cm 4cm", (1, 2, 3, 4)),
    ],
)
def test_a_margin_expands_the_way_css_says_it_does(written, centimetres):
    edges = page_margins(written)
    assert tuple(round(Emu(edge).cm, 2) for edge in edges) == centimetres


def test_a_bare_number_is_a_length_only_when_it_is_zero():
    assert length("0") == 0
    assert length("12") is None
    assert length("12pt") == Pt(12)


def test_the_sheet_comes_from_the_theme():
    section = convert(
        "Body.\n", theme=default_theme().with_page(size="Letter")
    ).sections[0]
    assert round(Emu(section.page_width).mm, 1) == round(LETTER_MM[0], 1)


def test_the_margins_come_from_the_theme():
    theme = default_theme().with_page(margin="3cm 2cm")
    section = convert("Body.\n", theme=theme).sections[0]
    assert round(Emu(section.top_margin).cm, 2) == 3
    assert round(Emu(section.left_margin).cm, 2) == 2


def test_geometry_word_cannot_read_warns_rather_than_failing():
    """The same theme makes a PDF perfectly well, so refusing would be worse."""
    theme = default_theme().with_page(size="A4 landscape portrait sideways")
    collected: list[str] = []
    render_docx(
        Document.from_markdown("Body.\n"),
        RenderOptions(theme=theme, warn=collected.append),
    )
    assert any("page size" in message for message in collected)


# --- the OOXML helpers ----------------------------------------------------


def test_a_properties_child_lands_in_the_order_the_schema_demands():
    """Word rejects the whole file when a child element is out of place."""
    document = read_docx()
    paragraph = document.add_paragraph("x")
    properties = paragraph._p.get_or_add_pPr()
    properties.get_or_add_ind()  # w:ind comes after w:shd
    properties.get_or_add_numPr()  # and w:numPr before it
    shade(properties, "#ff0000")
    assert [tag.rpartition("}")[2] for tag in (child.tag for child in properties)] == [
        "numPr",
        "shd",
        "ind",
    ]


def test_borders_are_written_with_the_edges_in_order():
    document = read_docx()
    properties = document.add_paragraph("x")._p.get_or_add_pPr()
    set_borders(properties, ["bottom", "top"], color="#d9d6de")
    borders = properties.find(qn("w:pBdr"))
    assert [child.tag.rpartition("}")[2] for child in borders] == ["top", "bottom"]
    assert borders[0].get(qn("w:color")) == "D9D6DE"


def test_a_heading_anchor_becomes_a_name_word_accepts():
    assert bookmark_name("inline-formatting") == "inline_formatting"
    assert bookmark_name("2026-in-review").startswith("_")
    assert len(bookmark_name("a" * 80)) == BOOKMARK_MAX_LENGTH


def test_an_external_link_becomes_a_relationship():
    document = read_docx()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("the project")
    link(paragraph, [run._r], url="https://example.com/x")
    element = paragraph._p.find(qn("w:hyperlink"))
    assert element is not None
    assert paragraph.part.rels[element.get(qn("r:id"))].target_ref == (
        "https://example.com/x"
    )
    assert element.find(qn("w:r")) is not None


def test_an_internal_link_points_at_a_bookmark_rather_than_a_url():
    document = read_docx()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("back up")
    link(paragraph, [run._r], anchor="a_heading")
    element = paragraph._p.find(qn("w:hyperlink"))
    assert element.get(qn("w:anchor")) == "a_heading"
    assert element.get(qn("r:id")) is None


def test_a_link_with_nothing_in_it_is_not_written_at_all():
    document = read_docx()
    paragraph = document.add_paragraph("plain")
    link(paragraph, [], url="https://example.com/x")
    assert paragraph._p.find(qn("w:hyperlink")) is None


def test_a_field_is_an_instruction_word_evaluates_when_it_opens_the_file():
    document = read_docx()
    paragraph = document.add_paragraph()
    field(paragraph, "PAGE", "1")
    xml = xml_of(paragraph._p)
    assert 'w:fldCharType="begin"' in xml
    assert '<w:instrText xml:space="preserve"> PAGE </w:instrText>' in xml
    assert 'w:fldCharType="end"' in xml


def test_each_numbering_instance_is_its_own_counter():
    document = read_docx()
    first = numbering_instance(document, "List Number")
    second = numbering_instance(document, "List Number", start=4)
    assert first != second

    numbering = document.part.numbering_part.element
    instances = {
        instance.get(qn("w:numId")): instance
        for instance in numbering.findall(qn("w:num"))
    }
    # Both draw the same bullets, and each starts counting where it was told.
    shapes = {
        str(first): instances[str(first)].find(qn("w:abstractNumId")).get(qn("w:val")),
        str(second): instances[str(second)]
        .find(qn("w:abstractNumId"))
        .get(qn("w:val")),
    }
    assert len(set(shapes.values())) == 1
    started = instances[str(second)].find(
        f"{qn('w:lvlOverride')}/{qn('w:startOverride')}"
    )
    assert started.get(qn("w:val")) == "4"


def test_a_header_row_is_marked_to_repeat_on_every_page():
    document = read_docx()
    table = document.add_table(rows=2, cols=2)
    repeat_as_header(table.rows[0])
    assert table.rows[0]._tr.find(f"{qn('w:trPr')}/{qn('w:tblHeader')}") is not None


def test_a_properties_child_is_not_added_twice():
    document = read_docx()
    properties = document.add_paragraph("x")._p.get_or_add_pPr()
    first = properties_child(properties, "w:shd")
    assert properties_child(properties, "w:shd") is first


# --- the walk -------------------------------------------------------------


def test_the_kitchen_sink_becomes_a_readable_document(kitchen_sink):
    document = reopen(render_docx(load_document(kitchen_sink), RenderOptions()).data)
    body = "\n".join(texts(document))
    for expected in [
        "Kitchen sink",
        "struck through",
        "Ordered second",
        "def convert",
        "attribution line",
        "A TOML file declaring",
        "Rendered by WeasyPrint",
    ]:
        assert expected in body, expected
    assert "Right column" in document.tables[0].rows[-1].cells[0].text


def test_headings_become_words_own_heading_styles():
    source = "".join(f"{'#' * level} Level {level}\n\n" for level in range(1, 7))
    assert styles(convert(source)) == [f"Heading {level}" for level in range(1, 7)]


def test_a_heading_deeper_than_word_goes_is_clamped():
    """Markdown stops at six and so does Word, but a tag is not a promise."""
    assert styles(convert("###### Six\n")) == ["Heading 6"]


def test_lists_use_words_list_styles_at_the_matching_level():
    document = convert("- one\n  - two\n    - three\n      - four\n")
    assert styles(document) == [
        "List Bullet",
        "List Bullet 2",
        "List Bullet 3",
        # Word defines three levels, so the fourth reuses the third.
        "List Bullet 3",
    ]


def test_ordered_lists_start_again_rather_than_carrying_on():
    """The defect a numbering instance per list exists to prevent.

    Word counts per numbering definition, not per list, so two ordered lists
    sharing the style ``List Number`` would number 1, 2, 3, 4, 5, 6.
    """
    document = convert("1. one\n2. two\n\nA paragraph.\n\n1. one again\n2. two again\n")
    numbered = [
        paragraph._p.find(f"{qn('w:pPr')}/{qn('w:numPr')}/{qn('w:numId')}")
        for paragraph in document.paragraphs
        if paragraph.style.name == "List Number"
    ]
    identifiers = [element.get(qn("w:val")) for element in numbered]
    assert identifiers[0] == identifiers[1]
    assert identifiers[2] == identifiers[3]
    assert identifiers[0] != identifiers[2]


def test_an_ordered_list_can_start_somewhere_other_than_one():
    document = convert("4. four\n5. five\n")
    element = document.paragraphs[0]._p.find(
        f"{qn('w:pPr')}/{qn('w:numPr')}/{qn('w:numId')}"
    )
    identifier = element.get(qn("w:val"))
    numbering = document.part.numbering_part.element
    instance = next(
        item
        for item in numbering.findall(qn("w:num"))
        if item.get(qn("w:numId")) == identifier
    )
    started = instance.find(f"{qn('w:lvlOverride')}/{qn('w:startOverride')}")
    assert started.get(qn("w:val")) == "4"


def test_a_task_list_is_printed_with_a_box_and_no_bullet():
    document = convert("- [ ] todo\n- [x] done\n")
    assert styles(document) == [BODY_STYLE, BODY_STYLE]
    assert texts(document) == ["☐ todo", "☑ done"]
    assert document.paragraphs[0]._p.find(f"{qn('w:pPr')}/{qn('w:numPr')}") is None


def test_inline_formatting_reaches_the_runs():
    document = convert("**b** *i* ~~s~~ `c`\n")
    formatted = {
        run.text: (run.bold, run.italic, run.font.strike, run.style.name)
        for run in document.paragraphs[0].runs
    }
    assert formatted["b"][0] is True
    assert formatted["i"][1] is True
    assert formatted["s"][2] is True
    assert formatted["c"][3] == CODE_INLINE_STYLE


def test_an_internal_link_and_its_heading_agree_on_the_bookmark():
    document = convert("# A Heading\n\nBack to [it](#a-heading).\n")
    marked = document.paragraphs[0]._p.find(qn("w:bookmarkStart"))
    anchor = document.paragraphs[1]._p.find(qn("w:hyperlink")).get(qn("w:anchor"))
    assert marked.get(qn("w:name")) == anchor


def test_a_local_image_is_embedded_and_fitted_to_the_column(kitchen_sink):
    document = reopen(render_docx(load_document(kitchen_sink), RenderOptions()).data)
    section = document.sections[0]
    column = (
        int(section.page_width) - int(section.left_margin) - int(section.right_margin)
    )
    assert len(document.inline_shapes) == 1
    assert document.inline_shapes[0].width <= column


def test_a_fenced_block_with_a_language_is_coloured():
    document = convert("```python\ndef f(x):\n    return x\n```\n")
    (block,) = [p for p in document.paragraphs if p.style.name == CODE_STYLE]
    colors = {str(run.font.color.rgb) for run in block.runs if run.font.color.rgb}
    assert len(block.runs) > 1
    assert colors
    assert block.text.startswith("def f(x):")


def test_a_fenced_block_with_no_language_is_one_plain_run():
    document = convert("```\n$ amethyst convert notes.md\n```\n")
    (block,) = [p for p in document.paragraphs if p.style.name == CODE_STYLE]
    (run,) = block.runs
    assert run.font.color.rgb is None


def test_highlighting_off_leaves_every_block_one_plain_run():
    document = convert("```python\ndef f(x): return x\n```\n", highlight_style="none")
    (block,) = [p for p in document.paragraphs if p.style.name == CODE_STYLE]
    (run,) = block.runs
    assert run.font.color.rgb is None


def test_a_dark_style_shades_every_code_block_not_only_the_coloured_ones():
    """A light box a paragraph away from a dark one is the drift to avoid."""
    document = convert(
        "```python\ndef f(): pass\n```\n\n```\nplain text\n```\n",
        highlight_style="monokai",
    )
    blocks = [p for p in document.paragraphs if p.style.name == CODE_STYLE]
    assert len(blocks) == 2
    for block in blocks:
        shading = block._p.get_or_add_pPr().find(qn("w:shd"))
        assert shading is not None
        assert shading.get(qn("w:fill")) == "272822"
        assert all(run.font.color.rgb is not None for run in block.runs)


def test_a_truncated_image_warns_rather_than_raising(tmp_path, fixtures):
    """A dropped download leaves half a file, so this path is now reachable."""
    whole = (fixtures / "assets" / "amethyst.png").read_bytes()
    (tmp_path / "half.png").write_bytes(whole[: len(whole) // 2])
    collected: list[str] = []
    render_docx(
        Document.from_markdown("![a](half.png)\n", base_dir=tmp_path),
        RenderOptions(warn=collected.append),
    )
    assert any("damaged" in message for message in collected)


def test_a_file_that_is_not_an_image_at_all_warns_too(tmp_path):
    (tmp_path / "nope.png").write_text("this is prose", encoding="utf-8")
    collected: list[str] = []
    render_docx(
        Document.from_markdown("![a](nope.png)\n", base_dir=tmp_path),
        RenderOptions(warn=collected.append),
    )
    assert any("image skipped" in message for message in collected)


def test_a_remote_image_is_reported_rather_than_fetched():
    """The walker never fetches: by here, an image is a local file or absent."""
    messages = warnings_from("![x](https://example.com/x.png)\n")
    assert any("not available locally" in message for message in messages)


def test_an_image_that_is_not_there_leaves_no_empty_paragraph(tmp_path):
    document = convert("![x](gone.png)\n", base_dir=tmp_path)
    assert texts(document) == []


def test_a_table_keeps_its_shape_alignment_and_header():
    document = convert(
        "| a | b |\n| - | :-: |\n| 1 | 2 |\n| 3 | 4 |\n",
    )
    table = document.tables[0]
    assert (len(table.rows), len(table.columns)) == (3, 2)
    assert [cell.text for cell in table.rows[0].cells] == ["a", "b"]
    assert table.rows[0]._tr.find(f"{qn('w:trPr')}/{qn('w:tblHeader')}") is not None
    assert table.cell(1, 1).paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert table.cell(0, 0).paragraphs[0].style.name == TABLE_TEXT_STYLE
    assert 'w:fill="F6F4F8"' in xml_of(table.cell(0, 0)._tc.get_or_add_tcPr())


def test_a_blockquote_is_styled_and_nesting_indents_it_further():
    document = convert("> outer\n>\n> > inner\n")
    assert styles(document) == [QUOTE_STYLE, QUOTE_STYLE]
    outer, inner = document.paragraphs
    assert outer.paragraph_format.left_indent > 0
    assert inner.paragraph_format.left_indent > outer.paragraph_format.left_indent


def test_quoting_indents_what_is_not_a_paragraph_too():
    """A list inside a quote has no Quote style to take the indent from."""
    document = convert("> 1. first\n>\n> ```\n> code\n> ```\n")
    listed, code = document.paragraphs
    assert listed.style.name == "List Number"
    assert code.style.name == CODE_STYLE
    assert listed.paragraph_format.left_indent > 0
    assert code.paragraph_format.left_indent > 0


def test_a_fenced_block_is_one_paragraph_that_keeps_its_indentation():
    document = convert("```python\ndef f():\n    return 1\n```\n")
    assert styles(document) == [CODE_STYLE]
    assert document.paragraphs[0].text == "def f():\n    return 1"


def test_a_horizontal_rule_is_an_empty_paragraph_with_a_border():
    document = convert("---\n")
    rule = document.paragraphs[0]
    assert rule.text == ""
    assert rule._p.find(f"{qn('w:pPr')}/{qn('w:pBdr')}/{qn('w:bottom')}") is not None


def test_a_definition_list_sets_the_term_apart_from_the_definition():
    document = convert("Theme\n: A TOML file.\n")
    term, definition = document.paragraphs
    assert all(run.bold for run in term.runs)
    assert definition.paragraph_format.left_indent > 0


def test_footnotes_become_an_endnote_list_with_superscript_references():
    document = convert("Text[^a].\n\n[^a]: The note.\n")
    reference = next(run for run in document.paragraphs[0].runs if run.font.superscript)
    assert reference.text == "1"
    assert document.paragraphs[-1].style.name == FOOTNOTE_STYLE
    assert document.paragraphs[-1].text == "1. The note."


def test_raw_html_is_skipped_with_one_warning_naming_its_line():
    messages = warnings_from("<div>\n  block\n</div>\n\nRaw <em>inline</em> too.\n")
    assert len(messages) == 2
    assert all("raw HTML" in message for message in messages)
    assert any("line 5" in message for message in messages)


def test_an_html_comment_is_skipped_without_a_word_about_it():
    """It was never going to be visible, so saying it was dropped is noise."""
    assert warnings_from("<!-- a note to self -->\n\nBody.\n") == []


def test_the_page_number_goes_in_the_footer_as_a_field():
    footer = convert("Body.\n").sections[0].footer
    assert "PAGE" in xml_of(footer.paragraphs[0]._p)


def test_page_numbers_can_be_suppressed():
    footer = convert("Body.\n", page_numbers=False).sections[0].footer
    assert "PAGE" not in xml_of(footer.paragraphs[0]._p)


def test_a_word_document_reports_no_page_count():
    """Word decides pagination when it opens the file, so there is none yet."""
    result = render_docx(Document.from_markdown("Body.\n"), RenderOptions())
    assert result.pages is None
    assert result.data.startswith(b"PK")


def test_the_file_does_not_claim_to_have_been_written_by_a_library():
    """python-docx's starter document says it was written in 2013, by itself."""
    properties = convert("Body.\n").core_properties
    assert properties.author == ""
    assert properties.comments == ""
    assert properties.created.year >= 2026


def test_the_frontmatter_becomes_the_properties_word_shows_in_its_info_pane():
    document = convert(
        "---\ntitle: T\nsubtitle: S\nauthor: A\ndate: 2026-08-31\n"
        "keywords: [one, two]\n---\n\nBody.\n"
    )
    properties = document.core_properties
    assert properties.title == "T"
    assert properties.author == "A"
    assert properties.subject == "S"
    assert properties.keywords == "one, two"
    # A declared date is the document's date, which is what "created" means.
    assert properties.created.date() == date(2026, 8, 31)


def test_a_date_word_could_not_read_leaves_the_timestamp_alone():
    properties = convert("---\ndate: Spring 2026\n---\n\nBody.\n").core_properties
    assert properties.created.year >= 2026


# --- the front matter -----------------------------------------------------


def test_a_title_page_is_built_from_the_frontmatter():
    document = convert(
        "---\ntitle: T\nsubtitle: S\nauthor: A\ndate: D\n---\n\n# H\n",
        title_page=True,
    )
    assert styles(document)[:4] == [
        TITLE_STYLE,
        SUBTITLE_STYLE,
        BODY_STYLE,
        COVER_DATE_STYLE,
    ]
    assert texts(document)[:4] == ["T", "S", "A", "D"]


def test_the_title_starts_down_the_page_as_the_stylesheet_pads_it_down():
    document = convert("---\ntitle: T\n---\n\n# H\n", title_page=True)
    theme = default_theme()
    assert document.paragraphs[0].paragraph_format.space_before == Pt(
        theme.type.size * TITLE_PAGE_OFFSET
    )


def test_the_author_sits_on_the_date_rather_than_a_paragraph_away_from_it():
    """The one thing `Normal` gets wrong on a cover: its gap is a body gap."""
    document = convert(
        "---\ntitle: T\nauthor: A\ndate: D\n---\n\n# H\n", title_page=True
    )
    author = document.paragraphs[1]
    assert author.text == "A"
    assert author.paragraph_format.space_after == Pt(
        default_theme().type.size * COVER_SPACE_AFTER
    )


def test_a_title_page_with_no_title_says_so_rather_than_printing_a_blank():
    messages = warnings_from("Body with no heading.\n", title_page=True)
    assert any("--title-page needs a title" in message for message in messages)


def test_the_contents_is_a_field_word_rebuilds_when_it_opens_the_file():
    document = convert("# One\n\n## Two\n", toc=True)
    xml = "".join(xml_of(paragraph._p) for paragraph in document.paragraphs)
    assert 'TOC \\o "1-2" \\h \\z \\u' in xml
    # Dirty, or Word shows the stored result until someone presses F9.
    assert 'w:fldChar w:fldCharType="begin" w:dirty="true"' in xml


def test_the_contents_entries_are_written_out_inside_the_field():
    """A reader that shows a field's stored result rather than evaluating it
    should still get a contents rather than a blank page."""
    document = convert("# One\n\n## Two\n\n### Three\n\n#### Four\n", toc=True)
    assert styles(document)[:5] == [
        TOC_HEADING_STYLE,
        "TOC 1",
        "TOC 2",
        "TOC 3",
        BODY_STYLE,  # the empty paragraph the section break is written on
    ]
    listed = texts(document)[: styles(document).index(BODY_STYLE)]
    assert listed[0] == CONTENTS_HEADING
    assert [entry.rstrip("\t") for entry in listed[1:]] == ["One", "Two", "Three"]


def test_an_entry_links_to_its_heading_and_asks_word_for_its_page():
    document = convert("## Inline formatting\n", toc=True)
    entry = xml_of(document.paragraphs[1]._p)
    name = bookmark_name("inline-formatting")
    assert f'w:anchor="{name}"' in entry
    assert f"PAGEREF {name}" in entry


def test_a_contents_with_nothing_to_list_says_so():
    messages = warnings_from("Body text.\n", toc=True)
    assert any("--toc needs headings" in message for message in messages)


def test_front_matter_becomes_a_section_of_its_own():
    """Which is what lets it carry different furniture, and what starts the
    document proper on a fresh page."""
    assert len(convert("# H\n", toc=True).sections) == 2
    assert len(convert("# H\n").sections) == 1


def test_the_front_matter_section_keeps_the_theme_page_geometry():
    """A section added after the theme was applied clones the one before it,
    but saying so out loud is cheaper than finding out it stopped."""
    theme = default_theme().with_page(size="Letter", margin="3cm")
    document = convert("---\ntitle: T\n---\n\n# H\n", title_page=True, theme=theme)
    assert len(document.sections) == 2
    for section in document.sections:
        assert round(Emu(section.page_width).mm, 1) == round(LETTER_MM[0], 1)
        assert round(Emu(section.left_margin).cm, 2) == 3


# --- the running head and the page number ---------------------------------


def test_the_head_names_the_title_and_the_section_word_is_in():
    document = convert("---\ntitle: T\n---\n\n# H\n\n## a\n\n## b\n")
    header = document.sections[-1].header
    assert not header.is_linked_to_previous
    xml = xml_of(header.paragraphs[0]._p)
    assert "<w:t>T</w:t>" in xml
    # Word's answer to the stylesheet's named string.
    assert 'STYLEREF "Heading 2"' in xml


def test_the_running_head_asks_nothing_of_the_reader_to_fill_itself_in():
    """Word works a STYLEREF out while it lays the page out, the way it works
    out a PAGE. Marking it dirty would buy nothing and cost an "update the
    fields in this document?" prompt on every open of every document."""
    document = convert("---\ntitle: T\n---\n\n# H\n\n## a\n\n## b\n")
    assert "w:dirty" not in xml_of(document.sections[-1].header.paragraphs[0]._p)
    assert "w:dirty" not in xml_of(document.sections[-1].footer.paragraphs[0]._p)


def test_only_a_contents_carries_a_field_word_has_to_be_asked_to_update():
    plain = convert("# H\n\n## a\n\n## b\n")
    listed = convert("# H\n\n## a\n\n## b\n", toc=True)
    assert "w:dirty" not in "".join(xml_of(p._p) for p in plain.paragraphs)
    assert "w:dirty" in "".join(xml_of(p._p) for p in listed.paragraphs)


def test_the_head_stops_at_the_title_when_no_level_repeats():
    document = convert("---\ntitle: T\n---\n\n# Only\n")
    xml = xml_of(document.sections[-1].header.paragraphs[0]._p)
    assert "<w:t>T</w:t>" in xml
    assert "STYLEREF" not in xml


def test_a_document_with_nothing_to_put_in_a_head_gets_none():
    document = convert("Body text.\n")
    assert document.sections[-1].header.is_linked_to_previous


def test_the_head_is_tabbed_to_the_edge_of_the_text_column_not_words_default():
    """The Header style's own tab stops sit where a US Letter sheet with
    one-inch margins puts them, and nowhere near the edge of any other."""
    document = convert("---\ntitle: T\n---\n\n# H\n\n## a\n\n## b\n")
    section = document.sections[-1]
    stops = document.sections[-1].header.paragraphs[0].paragraph_format.tab_stops
    assert [stop.position for stop in stops] == [
        section.page_width - section.left_margin - section.right_margin
    ]


def test_the_opening_page_of_a_plain_document_carries_no_head_but_is_numbered():
    """``@page :first`` by another name — Word has no such selector, so this
    is the "different first page" the format does have."""
    document = convert("---\ntitle: T\n---\n\n# H\n\n## a\n\n## b\n")
    section = document.sections[0]
    assert section.different_first_page_header_footer
    assert not section.first_page_footer.is_linked_to_previous
    assert "PAGE" in xml_of(section.first_page_footer.paragraphs[0]._p)


def test_a_cover_is_not_numbered_and_the_contents_behind_it_is():
    document = convert("---\ntitle: T\n---\n\n# H\n", toc=True, title_page=True)
    front = document.sections[0]
    assert front.different_first_page_header_footer
    assert front.first_page_footer.paragraphs[0].text == ""
    assert "PAGE" in xml_of(front.footer.paragraphs[0]._p)


def test_a_contents_with_no_cover_in_front_of_it_is_numbered_from_its_first_page():
    front = convert("# H\n", toc=True).sections[0]
    assert not front.different_first_page_header_footer


def test_no_page_numbers_means_no_footer_at_all():
    """A page with nothing at the foot of it, rather than an empty line."""
    document = convert("Body.\n", page_numbers=False)
    assert document.sections[0].footer.is_linked_to_previous


def test_the_paragraph_after_a_table_is_given_the_gap_word_leaves_out():
    document = convert("| a |\n| - |\n| 1 |\n\nAfter.\n")
    assert document.paragraphs[-1].paragraph_format.space_before > 0


def test_a_heading_after_a_table_keeps_its_own_larger_gap():
    document = convert("| a |\n| - |\n| 1 |\n\n## After\n")
    heading = document.paragraphs[-1]
    assert heading.paragraph_format.space_before is None


def test_a_character_style_is_defined_for_code_and_for_links():
    document = blank_docx()
    for name in (CODE_INLINE_STYLE, "Hyperlink"):
        assert document.styles[name].type == WD_STYLE_TYPE.CHARACTER
