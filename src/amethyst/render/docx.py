"""The token stream, walked into a Word document.

Word has no equivalent of CSS paged media, so there is nothing here to hand
the job to: a DOCX file is a flat sequence of styled paragraphs and runs, and
this module is the walk that produces it. That is why the two pipelines are
separate. What keeps their output the same document is the theme — every style
this module names was defined by :mod:`amethyst.theme.to_docx` from the same
declaration the stylesheet was generated from, and nothing below chooses a
font, a size or a colour.

Three things about the walk itself.

markdown-it's stream is flat, with ``_open`` and ``_close`` tokens marking
nesting, so containers are handled by finding the matching close and recursing
over the span between. The alternative — a state machine over a flat loop —
is the same program with the structure hidden.

Where a paragraph goes is context, not content. A paragraph inside a list item
carries the item's bullet; inside a blockquote it is indented and set in the
quoted style; inside a footnote it is small and muted. All of that is decided
in one place, :meth:`_Builder._paragraph`, so no handler has to know what it
is nested inside.

And several things Markdown says have no Word equivalent at all. Raw HTML is
skipped with a warning naming its line, footnotes become an endnote-style list
at the end rather than real Word footnotes, and a task list gets a printed
checkbox rather than a real one. Each is a deliberate approximation from the
feature matrix, not an oversight.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, replace
from datetime import datetime, time, timezone
from pathlib import Path
from typing import Any
from urllib.parse import unquote, urlsplit

from docx import Document as new_docx
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.image.exceptions import (
    InvalidImageStreamError,
    UnexpectedEndOfFileError,
    UnrecognizedImageError,
)
from docx.oxml.ns import qn
from docx.shared import Emu, Length, Pt
from markdown_it.token import Token

from amethyst.document import Document
from amethyst.errors import RenderError
from amethyst.ooxml import (
    bookmark,
    bookmark_name,
    field,
    field_end,
    field_start,
    link,
    numbering_instance,
    repeat_as_header,
    set_borders,
    set_numbering,
    shade,
)
from amethyst.parse.assets import REMOTE_SCHEMES
from amethyst.render.base import RenderOptions, RenderResult
from amethyst.render.furniture import (
    CONTENTS_HEADING,
    contents,
    cover,
    outline_depth,
    section_level,
)
from amethyst.render.highlight import Highlighter, Span
from amethyst.theme.to_docx import (
    BODY_STYLE,
    BULLET_STYLES,
    CODE_BOX_PADDING,
    CODE_BOX_SIDES,
    CODE_INLINE_STYLE,
    CODE_STYLE,
    COVER_DATE_STYLE,
    COVER_SPACE_AFTER,
    FOOTER_STYLE,
    FOOTNOTE_STYLE,
    HEADER_STYLE,
    HEADING_STYLES,
    LINK_STYLE,
    LIST_INDENT_STEP,
    NUMBER_STYLES,
    QUOTE_STYLE,
    SUBTITLE_STYLE,
    TABLE_STYLE,
    TABLE_TEXT_STYLE,
    TITLE_PAGE_OFFSET,
    TITLE_STYLE,
    TOC_HEADING_STYLE,
    TOC_STYLES,
    apply_page,
    apply_theme,
    quote_indent,
    rgb,
    text_width,
)

#: python-docx's three ways of saying "that is not a picture I can embed".
#: They share no base class, so all three have to be named — and all three are
#: reachable from a document: a file that is not an image, one that is an image
#: format Word has no part type for, and one that was cut off. The last is why
#: this matters more since images are downloaded: a truncated file is what a
#: dropped connection leaves behind.
UNUSABLE_IMAGE = (
    InvalidImageStreamError,
    UnexpectedEndOfFileError,
    UnrecognizedImageError,
)

#: What a task list item is printed as. Word has real checkbox controls, but
#: they are form fields tied to a content control, and a document that has to
#: be filled in is not what a converted Markdown list is.
CHECKED = "☑"
UNCHECKED = "☐"

#: The instructions Word evaluates for itself: the number of the page a field
#: lands on, the page a bookmark is on, the text of the nearest heading of one
#: style — Word's answer to a CSS named string — and the table of contents
#: itself, built from heading levels 1 to N.
#:
#: The first and third are worked out while Word lays the pages out and need
#: nothing asked of the reader. The other two are not, and are marked dirty so
#: that Word fills them in on open — which is what raises its "update the
#: fields in this document?" prompt, and why only ``--toc`` raises it.
PAGE_FIELD = "PAGE"
PAGE_REFERENCE_FIELD = "PAGEREF {name} \\h"
SECTION_FIELD = 'STYLEREF "{style}" \\* MERGEFORMAT'
CONTENTS_FIELD = 'TOC \\o "1-{depth}" \\h \\z \\u'

#: The air around a horizontal rule, and above the line that opens the
#: footnotes, as multiples of the body size. Both mirror the margins the
#: structural stylesheet gives the same two elements, halved for the rule
#: because CSS collapses a margin against its neighbour's and Word adds them.
RULE_GAP = 0.8
FOOTNOTES_GAP = 2.5

#: Cell alignment as markdown-it writes it: an inline style on the cell.
ALIGNMENTS = {
    "text-align:left": WD_ALIGN_PARAGRAPH.LEFT,
    "text-align:center": WD_ALIGN_PARAGRAPH.CENTER,
    "text-align:right": WD_ALIGN_PARAGRAPH.RIGHT,
}

#: An HTML block that is nothing but a comment. It was never going to be
#: visible in either format, so warning that it was skipped is noise.
HTML_COMMENT = re.compile(r"\A\s*(?:<!--.*?-->\s*)+\Z", re.DOTALL)

#: The checkbox the tasklist plugin emits. It arrives as raw HTML inside the
#: item's inline token rather than as a token of its own, so without this the
#: rule that skips raw HTML would quietly eat every checkbox in the document.
TASK_CHECKBOX = re.compile(r"<input[^>]*\bclass=\"task-list-item-checkbox\"", re.I)


def render_docx(document: Document, options: RenderOptions) -> RenderResult:
    """Convert a document to the bytes of a Word file."""
    return _Builder(document, options).build()


@dataclass(frozen=True)
class _Run:
    """The character formatting in force at one point in an inline walk."""

    bold: bool = False
    italic: bool = False
    strike: bool = False
    code: bool = False
    link: bool = False
    superscript: bool = False


#: No emphasis, no link, no code: what an inline walk starts from unless the
#: block it sits in says otherwise.
_PLAIN = _Run()


@dataclass(frozen=True)
class _ListLevel:
    """One open list, and how the items inside it are marked and indented."""

    style: str
    #: The list's own counter, or ``None`` for a list that is not numbered.
    number: int | None
    indent: Length


class _Builder:
    """One conversion. Not reused: the state below belongs to one document."""

    def __init__(self, document: Document, options: RenderOptions) -> None:
        """Start an empty document with the theme already compiled into it."""
        self._document = document
        self._options = options
        self._theme = options.theme
        self._warn = options.warn
        self._docx = new_docx()
        self._highlighter = Highlighter(options.highlight_style, warn=self._warn)
        apply_theme(self._docx, self._theme, warn=self._warn)

        self._lists: list[_ListLevel] = []
        self._pending: _ListLevel | None = None
        self._quotes = 0
        self._footnotes = 0
        self._prefix: str | None = None
        self._bookmarks = 0
        self._after_table = False
        self._extra_indent: Length | None = None
        self._warned_html: set[int | None] = set()

    # --- the document ------------------------------------------------------

    def build(self) -> RenderResult:
        """Walk the whole document and return the file's bytes."""
        self._properties()
        if self._front_matter():
            # The front matter becomes a section of its own so that it can
            # carry different furniture from the body — no running head, and
            # no page number on a cover. It is also what starts the document
            # proper on a fresh page, which is the break the stylesheet gets
            # from `break-after: page`.
            self._docx.add_section(WD_SECTION.NEW_PAGE)
            apply_page(self._docx.sections[-1], self._theme, warn=self._warn)
        self._blocks(self._document.tokens, 0, len(self._document.tokens))
        self._furniture()
        return RenderResult(data=self._save(), pages=None)

    def _properties(self) -> None:
        """Say what the document says about itself, and nothing else.

        python-docx builds every document from a template whose properties
        claim it was written by "python-docx" in 2013, and Word shows those in
        the file's info pane. They are replaced by the frontmatter — the same
        four fields the PDF carries as its metadata — so that a reader who
        opens the information pane sees the same thing in either format.
        """
        properties = self._docx.core_properties
        properties.author = self._document.author or ""
        properties.title = self._document.title or ""
        properties.subject = self._document.subtitle or ""
        properties.keywords = self._document.keywords or ""
        properties.last_modified_by = ""
        properties.comments = ""
        properties.revision = 1
        now = datetime.now(tz=timezone.utc)
        declared = self._document.created
        # A declared date is the document's date, which is what the format
        # means by "created". A date it could not read — "Spring 2026" — stays
        # on the title page and out of the timestamp.
        properties.created = (
            datetime.combine(declared, time(), tzinfo=timezone.utc)
            if declared is not None
            else now
        )
        properties.modified = now

    def _save(self) -> bytes:
        """Serialise the finished document, without ever touching the disk."""
        buffer = io.BytesIO()
        try:
            self._docx.save(buffer)
        except OSError as exc:  # pragma: no cover - an in-memory write
            raise RenderError(f"Could not build the Word document: {exc}.") from exc
        return buffer.getvalue()

    # --- front matter ------------------------------------------------------

    def _front_matter(self) -> bool:
        """Write the cover and the contents, and say whether either happened."""
        written = False
        if self._options.title_page:
            written |= self._title_page()
        if self._options.toc:
            written |= self._contents()
        return written

    def _title_page(self) -> bool:
        """A cover built from the frontmatter, in the styles the theme set."""
        page = cover(self._document)
        if page is None:
            self._warn(
                "--title-page needs a title; the document declares none, so "
                "no title page was made."
            )
            return False
        title = self._docx.add_paragraph(style=self._docx.styles[TITLE_STYLE])
        # The stylesheet pads the cover down the sheet; Word has no padding, so
        # the same gap is set above the one paragraph it would have pushed.
        title.paragraph_format.space_before = Pt(
            self._theme.type.size * TITLE_PAGE_OFFSET
        )
        title.add_run(page.title)
        for style, value in (
            (SUBTITLE_STYLE, page.subtitle),
            (BODY_STYLE, page.author),
            (COVER_DATE_STYLE, page.date),
        ):
            if value:
                line = self._docx.add_paragraph(value, style=self._docx.styles[style])
                if style == BODY_STYLE:
                    # The author sits on the date rather than a paragraph's gap
                    # away from it, which is the one thing `Normal` gets wrong
                    # on a cover.
                    line.paragraph_format.space_after = Pt(
                        self._theme.type.size * COVER_SPACE_AFTER
                    )
        return True

    def _contents(self) -> bool:
        """The contents, as a TOC field whose result is already filled in.

        The field is what makes this a real Word table of contents: it is
        marked dirty, so Word rebuilds it against its own pagination the
        moment the file opens, and it stays right when the document is edited.
        Writing the entries out inside it as well costs little and means every
        other reader — one that shows a field's stored result rather than
        evaluating it — has a contents rather than a blank page.
        """
        entries = contents(self._document, self._options.toc_depth)
        if not entries:
            self._warn(
                "--toc needs headings; the document has none, so no contents was made."
            )
            return False

        heading = self._docx.add_paragraph(style=self._docx.styles[TOC_HEADING_STYLE])
        heading.add_run(CONTENTS_HEADING)

        instruction = CONTENTS_FIELD.format(
            depth=outline_depth(entries, self._options.toc_depth)
        )
        paragraph = None
        for index, entry in enumerate(entries):
            level = min(entry.level, len(TOC_STYLES))
            paragraph = self._docx.add_paragraph(
                style=self._docx.styles[TOC_STYLES[level - 1]]
            )
            if index == 0:
                for element in field_start(instruction, dirty=True):
                    paragraph._p.append(element)
            self._contents_entry(paragraph, entry.text, entry.anchor)
        if paragraph is not None:
            paragraph._p.append(field_end())
        return True

    def _contents_entry(self, paragraph: Any, text: str, anchor: str | None) -> None:
        """One line of the contents: the heading, a leader, and its page.

        An entry with no anchor gets neither the link nor the number, because
        both are the same bookmark by two names. Listing it unlinked beats
        dropping a heading out of the contents without saying so.
        """
        run = paragraph.add_run(text)
        if anchor is None:
            return
        name = bookmark_name(anchor)
        link(paragraph, [run._r], anchor=name)
        paragraph.add_run().add_tab()
        # Dirty, because the placeholder is empty: nothing here knows which
        # page a heading will land on, and Word does as soon as it opens.
        field(paragraph, PAGE_REFERENCE_FIELD.format(name=name), dirty=True)

    # --- page furniture ----------------------------------------------------

    def _furniture(self) -> None:
        """Put the running head and the page number on the pages that get them.

        The two formats are made to agree here, and the agreement is worth
        stating: the opening page of a document carries no head, whatever is
        on it; a cover carries no page number either; and the front matter,
        which belongs to no section, carries no head at all. In CSS that is
        ``@page :first`` and a named page. In Word there is no such selector,
        so it is a section for the front matter, and Word's own "different
        first page" for a document that has none.
        """
        sections = self._docx.sections
        front = sections[0] if len(sections) > 1 else None
        body = sections[-1]
        head = self._running_head()

        if front is not None:
            # A section with no header part simply has none, which is what the
            # front matter wants — so the only thing to say is that the cover
            # is not to be numbered.
            front.different_first_page_header_footer = self._options.title_page
            if self._options.title_page:
                # Explicit rather than inherited: a section marked "different
                # first page" with nothing defined for it is empty by
                # inheritance, and inheriting from nothing is a fact about the
                # format that is cheaper to state than to rely on.
                front.first_page_footer.is_linked_to_previous = False
            self._number(front.footer)
            body.different_first_page_header_footer = False
        else:
            # Nothing precedes the body, so its first page is the document's
            # opening page and takes the head off in the only way Word offers.
            body.different_first_page_header_footer = head is not None
            if head is not None:
                self._number(body.first_page_footer)
        if head is not None:
            self._head(body.header, head)
        self._number(body.footer)

    def _running_head(self) -> tuple[str | None, int | None] | None:
        """What the head says: the title, and the level it tracks. Or nothing."""
        title = self._document.title
        level = section_level(self._document)
        if not title and level is None:
            return None
        return (title, level)

    def _head(self, header: Any, head: tuple[str | None, int | None]) -> None:
        """Write the running head: the title left, the current section right.

        Word's answer to the stylesheet's named string is ``STYLEREF``, which
        names the nearest heading of one style. The tab stop is set here rather
        than left to the ``Header`` style's own, which sits where a US Letter
        sheet with one-inch margins puts it and nowhere near the edge of any
        other column.
        """
        title, level = head
        header.is_linked_to_previous = False
        paragraph = header.paragraphs[0]
        paragraph.style = self._docx.styles[HEADER_STYLE]
        width = text_width(self._docx.sections[-1])
        if width is not None:
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                width, WD_TAB_ALIGNMENT.RIGHT
            )
        if title:
            paragraph.add_run(title)
        if level is not None:
            paragraph.add_run().add_tab()
            # Not marked dirty, and deliberately: Word works a STYLEREF out
            # while it lays the page out, the same way it works out a PAGE, so
            # the head fills itself in with nothing asked of the reader.
            # Marking it would cost a "do you want to update the fields in this
            # document?" on every open and buy nothing. Verified in Word.
            field(paragraph, SECTION_FIELD.format(style=HEADING_STYLES[level - 1]))

    def _number(self, footer: Any) -> None:
        """Put the page number in a footer, as the PDF puts it in the margin.

        A field rather than a number: nothing here knows how many pages Word
        will decide the document has, and a field is how the format says "the
        number of the page this lands on". With ``--no-page-numbers`` no
        footer is defined at all, which is a page with nothing at the foot of
        it rather than a page with an empty line there.
        """
        if not self._options.page_numbers:
            return
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.style = self._docx.styles[FOOTER_STYLE]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        field(paragraph, PAGE_FIELD, "1")

    # --- blocks ------------------------------------------------------------

    def _blocks(self, tokens: list[Token], start: int, end: int) -> None:
        """Walk a run of block tokens, in order, until the end of the range."""
        index = start
        while index < end:
            index = self._block(tokens, index, end)

    def _block(self, tokens: list[Token], index: int, end: int) -> int:
        """Handle one block, and return the index of the next one."""
        token = tokens[index]
        handler = _BLOCKS.get(token.type)
        if handler is None:
            # Everything the walk does not act on is either a close tag whose
            # open tag consumed the span, or a token that carries no content.
            return index + 1
        return handler(self, tokens, index, end)

    def _heading(self, tokens: list[Token], index: int, end: int) -> int:
        """A heading, styled by level and bookmarked so the contents can link."""
        token = tokens[index]
        level = min(int(token.tag[1:]), len(HEADING_STYLES))
        paragraph = self._paragraph(HEADING_STYLES[level - 1])
        anchor = token.attrGet("id")
        if isinstance(anchor, str) and anchor:
            self._bookmarks += 1
            bookmark(paragraph, bookmark_name(anchor), self._bookmarks)
        close = _closing(tokens, index, end)
        self._inline_span(paragraph, tokens, index + 1, close)
        return close + 1

    def _paragraph_block(self, tokens: list[Token], index: int, end: int) -> int:
        """An ordinary paragraph, dropped again if nothing survived the walk."""
        close = _closing(tokens, index, end)
        paragraph = self._paragraph()
        self._inline_span(paragraph, tokens, index + 1, close)
        # A paragraph whose whole content was dropped — one holding nothing
        # but an image that could not be loaded, say — would otherwise be a
        # blank line where the author wrote something.
        if not paragraph.runs and paragraph._p.find(qn("w:hyperlink")) is None:
            paragraph._p.getparent().remove(paragraph._p)
        return close + 1

    def _code(self, tokens: list[Token], index: int, _end: int) -> int:
        """A fenced block: one shaded paragraph, highlighted run by run."""
        token = tokens[index]
        paragraph = self._paragraph(CODE_STYLE)
        background = self._highlighter.background
        if background is not None:
            # A dark highlighting style is a panel of its own, and every code
            # block in the document is one — including the blocks it could not
            # colour, which would otherwise be the theme's light box a
            # paragraph away from a dark one. The stylesheet says the same
            # thing with a bare `pre` rule.
            properties = paragraph._p.get_or_add_pPr()
            shade(properties, background)
            set_borders(
                properties,
                CODE_BOX_SIDES,
                color=background,
                space=CODE_BOX_PADDING,
            )
        # `info` is what was written after the fence — the language, and
        # anything else on the line, which nothing here reads. An indented
        # block has no info at all, and so is never highlighted.
        language = (token.info or "").split(maxsplit=1)
        spans = self._highlighter.spans(token.content, language[0] if language else "")
        if spans is None:
            spans = [
                Span(
                    text=token.content.rstrip("\n"),
                    color=self._highlighter.foreground,
                )
            ]
        for span in spans:
            self._span(paragraph, span)
        return index + 1

    def _span(self, paragraph: Any, span: Span) -> None:
        """One coloured run of code, inside the paragraph holding the block."""
        run = paragraph.add_run()
        # The setter turns each newline into a line break rather than a new
        # paragraph, which keeps one fenced block inside one shaded box.
        run.text = span.text
        if span.color is not None:
            run.font.color.rgb = rgb(span.color)
        if span.bold:
            run.bold = True
        if span.italic:
            run.italic = True

    def _quote(self, tokens: list[Token], index: int, end: int) -> int:
        """A blockquote, which is a depth rather than a container: what it holds
        is ordinary blocks, indented by how many quotes are open around them."""
        close = _closing(tokens, index, end)
        self._quotes += 1
        self._blocks(tokens, index + 1, close)
        self._quotes -= 1
        return close + 1

    def _list(self, tokens: list[Token], index: int, end: int) -> int:
        """A list, given a counter of its own so it does not continue the last."""
        token = tokens[index]
        close = _closing(tokens, index, end)
        ordered = token.type == "ordered_list_open"
        # A task list carries its own marker in the checkbox, so it gets the
        # indent of a list and none of the bullets — which is exactly what the
        # stylesheet does with `.contains-task-list` on the PDF side.
        tasks = "task-list" in str(token.attrGet("class") or "")
        level = min(len(self._lists), len(BULLET_STYLES) - 1)
        indent = Emu(int(LIST_INDENT_STEP) * (level + 1))

        if tasks:
            entry = _ListLevel(style=BODY_STYLE, number=None, indent=indent)
        else:
            style = (NUMBER_STYLES if ordered else BULLET_STYLES)[level]
            entry = _ListLevel(
                style=style,
                number=numbering_instance(self._docx, style, start=_start(token)),
                indent=indent,
            )

        self._lists.append(entry)
        self._blocks(tokens, index + 1, close)
        self._lists.pop()
        return close + 1

    def _item(self, tokens: list[Token], index: int, end: int) -> int:
        """One item, whose marker goes on the first paragraph inside it."""
        close = _closing(tokens, index, end)
        self._pending = self._lists[-1] if self._lists else None
        self._blocks(tokens, index + 1, close)
        # An item whose only content was something that starts no paragraph
        # would otherwise leak its marker onto whatever comes next.
        self._pending = None
        return close + 1

    def _rule(self, _tokens: list[Token], index: int, _end: int) -> int:
        """A horizontal rule: an empty paragraph with a border underneath it."""
        paragraph = self._paragraph()
        set_borders(
            paragraph._p.get_or_add_pPr(), ["bottom"], color=self._theme.colors.rule
        )
        gap = Pt(self._theme.type.size * RULE_GAP)
        paragraph.paragraph_format.space_before = gap
        paragraph.paragraph_format.space_after = gap
        return index + 1

    def _term(self, tokens: list[Token], index: int, end: int) -> int:
        """The term of a definition list, set bold above its definition."""
        close = _closing(tokens, index, end)
        paragraph = self._paragraph()
        self._inline_span(paragraph, tokens, index + 1, close, base=_Run(bold=True))
        return close + 1

    def _definition(self, tokens: list[Token], index: int, end: int) -> int:
        """A definition is blocks of its own, set in under the term above it."""
        close = _closing(tokens, index, end)
        outer = self._extra_indent
        self._extra_indent = Pt(self._theme.type.size * self._theme.spacing.indent)
        self._blocks(tokens, index + 1, close)
        self._extra_indent = outer
        return close + 1

    def _html(self, tokens: list[Token], index: int, _end: int) -> int:
        """A raw HTML block, which Word has no way to hold: skip it and warn."""
        self._skip_html(tokens[index])
        return index + 1

    def _footnote_block(self, tokens: list[Token], index: int, end: int) -> int:
        """Open the endnote-style list the footnote bodies are collected into."""
        close = _closing(tokens, index, end)
        separator = self._paragraph(FOOTNOTE_STYLE)
        set_borders(
            separator._p.get_or_add_pPr(), ["bottom"], color=self._theme.colors.rule
        )
        separator.paragraph_format.space_before = Pt(
            self._theme.type.size * FOOTNOTES_GAP
        )
        self._footnotes += 1
        self._blocks(tokens, index + 1, close)
        self._footnotes -= 1
        return close + 1

    def _footnote(self, tokens: list[Token], index: int, end: int) -> int:
        """One footnote body, numbered to match the reference that points at it."""
        close = _closing(tokens, index, end)
        self._prefix = f"{_footnote_number(tokens[index])}. "
        self._blocks(tokens, index + 1, close)
        self._prefix = None
        return close + 1

    # --- tables ------------------------------------------------------------

    def _table(self, tokens: list[Token], index: int, end: int) -> int:
        """A GFM table, sized to its widest row and given a repeating header."""
        close = _closing(tokens, index, end)
        rows = _table_rows(tokens, index + 1, close)
        if not rows:
            return close + 1

        width = max(len(cells) for _, cells in rows)
        table = self._docx.add_table(rows=len(rows), cols=width, style=TABLE_STYLE)
        table.autofit = True
        set_borders(
            table._tbl.tblPr,
            ["top", "left", "bottom", "right", "insideH", "insideV"],
            color=self._theme.colors.rule,
        )

        for row_index, (header, cells) in enumerate(rows):
            if header:
                repeat_as_header(table.rows[row_index])
            for column, (content, alignment) in enumerate(cells):
                self._cell(table.cell(row_index, column), content, alignment, header)

        # Word puts nothing between a table and what follows it, and this is
        # cheaper than the empty paragraph the usual workaround adds.
        self._after_table = True
        return close + 1

    def _cell(
        self, cell: Any, content: Token | None, alignment: str | None, header: bool
    ) -> None:
        """One cell: the column's alignment, and a fill if it is a header."""
        paragraph = cell.paragraphs[0]
        paragraph.style = self._docx.styles[TABLE_TEXT_STYLE]
        if alignment in ALIGNMENTS:
            paragraph.alignment = ALIGNMENTS[alignment]
        if header:
            shade(cell._tc.get_or_add_tcPr(), self._theme.colors.fill)
        if content is not None:
            self._inline(paragraph, content, base=_Run(bold=header))

    # --- inline ------------------------------------------------------------

    def _inline_span(
        self,
        paragraph: Any,
        tokens: list[Token],
        start: int,
        end: int,
        *,
        base: _Run = _PLAIN,
    ) -> None:
        """Render whichever of the tokens in a span carries the inline text."""
        for index in range(start, end):
            if tokens[index].type == "inline":
                self._inline(paragraph, tokens[index], base=base)

    def _inline(self, paragraph: Any, token: Token, *, base: _Run = _PLAIN) -> None:
        """Walk one inline token into runs, carrying the formatting as it opens
        and closes. Links are the awkward case: their runs are written into the
        paragraph first and moved inside the hyperlink element on close."""
        style = base
        line = token.map[0] + 1 if token.map else None
        # Where each open link points, and the runs written since it opened,
        # so that they can be moved inside it when it closes. The target is
        # kept from the opening token because the closing one does not carry
        # it — an easy thing to get wrong, and silent when you do: the runs
        # come out styled as a link that goes nowhere.
        collected: list[tuple[str, list[Any]]] = []

        for child in token.children or []:
            kind = child.type
            if kind == "text":
                # markdown-it leaves empty text tokens behind where it merged
                # adjacent ones, and a run with nothing in it is still a run.
                if child.content:
                    self._add(paragraph, child.content, style, collected)
            elif kind == "softbreak":
                self._add(paragraph, " ", style, collected)
            elif kind == "hardbreak":
                self._add(paragraph, "", style, collected).add_break()
            elif kind == "code_inline":
                self._add(
                    paragraph, child.content, replace(style, code=True), collected
                )
            elif kind in _OPENS:
                style = replace(style, **{_OPENS[kind]: True})
            elif kind in _CLOSES:
                style = replace(style, **{_CLOSES[kind]: False})
            elif kind == "link_open":
                style = replace(style, link=True)
                collected.append((_href(child), []))
            elif kind == "link_close":
                style = replace(style, link=False)
                self._close_link(paragraph, collected)
            elif kind == "image":
                self._image(paragraph, child, line, collected)
            elif kind == "footnote_ref":
                number = str(_footnote_number(child))
                self._add(
                    paragraph, number, replace(style, superscript=True), collected
                )
            elif kind == "html_inline":
                self._inline_html(paragraph, child, style, collected, line)
            elif kind == "footnote_anchor":
                # The back-reference to where the footnote was cited. The PDF
                # hides it too: it is a browser affordance, and on paper the
                # link it offers goes nowhere.
                continue

    def _add(
        self,
        paragraph: Any,
        text: str,
        style: _Run,
        collected: list[tuple[str, list[Any]]],
    ) -> Any:
        """Add one run, formatted as the walk currently says, and record it."""
        run = paragraph.add_run(text)
        if style.code:
            run.style = self._docx.styles[CODE_INLINE_STYLE]
            if style.link:
                # A run carries one character style, so a link that is also
                # code keeps the code style and takes the link's colour.
                run.font.color.rgb = rgb(self._theme.colors.accent)
        elif style.link:
            run.style = self._docx.styles[LINK_STYLE]
        if style.bold:
            run.bold = True
        if style.italic:
            run.italic = True
        if style.strike:
            run.font.strike = True
        if style.superscript:
            run.font.superscript = True
        for _, pending in collected:
            pending.append(run._r)
        return run

    def _close_link(
        self, paragraph: Any, collected: list[tuple[str, list[Any]]]
    ) -> None:
        """Wrap the runs written since a link opened in the link itself."""
        if not collected:
            return
        href, elements = collected.pop()
        # Links are recorded but never rewritten by asset resolution, so this
        # is the author's own reference: a URL, or a fragment naming a heading
        # in this document, which is a bookmark by the time it gets here.
        if href.startswith("#"):
            link(paragraph, elements, anchor=bookmark_name(href[1:]))
        elif href:
            link(paragraph, elements, url=href)

    def _image(
        self,
        paragraph: Any,
        token: Token,
        line: int | None,
        collected: list[tuple[str, list[Any]]],
    ) -> None:
        """An inline image, scaled to the text column, or a warning naming it."""
        source = token.attrGet("src")
        source = source if isinstance(source, str) else ""
        where = f" (line {line})" if line is not None else ""
        if urlsplit(source).scheme in REMOTE_SCHEMES:
            self._warn(f"remote image not available locally: {source}{where}")
            return

        path = Path(unquote(urlsplit(source).path))
        if not path.is_absolute():
            path = self._document.base_dir / path
        run = paragraph.add_run()
        try:
            picture = run.add_picture(str(path))
        except (OSError, ValueError, *UNUSABLE_IMAGE) as exc:
            self._warn(f"image skipped: {source}{where} — {_why_unusable(exc)}.")
            run._r.getparent().remove(run._r)
            return
        _fit(picture, self._column_width())
        for _, pending in collected:
            pending.append(run._r)

    def _inline_html(
        self,
        paragraph: Any,
        token: Token,
        style: _Run,
        collected: list[tuple[str, list[Any]]],
        line: int | None,
    ) -> None:
        """Raw HTML inside a paragraph — which is where a checkbox arrives."""
        if TASK_CHECKBOX.search(token.content):
            checked = "checked" in token.content.lower()
            self._add(paragraph, CHECKED if checked else UNCHECKED, style, collected)
            return
        # An inline token carries no line map of its own; the paragraph it sits
        # in does, and naming that line is what makes the warning actionable.
        self._skip_html(token, line)

    # --- shared ------------------------------------------------------------

    def _paragraph(self, style: str = BODY_STYLE) -> Any:
        """Start a paragraph, in whatever the walk is currently inside.

        This is the one place that knows what nesting means, so that no
        handler above has to: a list item's marker, a blockquote's indent and
        style, a footnote's smaller type and its number, and the gap Word
        would otherwise not leave under a table all land here.
        """
        item, self._pending = self._pending, None
        prefix, self._prefix = self._prefix, None
        plain = style == BODY_STYLE

        if plain and item is not None:
            style = item.style
        elif plain and self._footnotes:
            style = FOOTNOTE_STYLE
        elif plain and self._quotes:
            style = QUOTE_STYLE

        paragraph = self._docx.add_paragraph(style=self._docx.styles[style])
        if item is not None and item.number is not None:
            set_numbering(paragraph, item.number)
        indent = self._indent(item)
        if indent:
            paragraph.paragraph_format.left_indent = indent
        if self._after_table:
            self._open_after_table(paragraph, style)
            self._after_table = False
        if prefix:
            paragraph.add_run(prefix)
        return paragraph

    def _open_after_table(self, paragraph: Any, style: str) -> None:
        """Leave the gap Word does not leave of its own accord after a table.

        A table has no space below it, so whatever follows sits flush against
        its bottom rule. The gap belongs on the paragraph that follows rather
        than on an empty paragraph of its own, which would be a blank line —
        and a style that already asks for more room than that keeps its own.
        """
        gap = Pt(self._theme.type.size * self._theme.spacing.block)
        declared = self._docx.styles[style].paragraph_format.space_before
        if declared is None or declared < gap:
            paragraph.paragraph_format.space_before = gap

    def _indent(self, item: _ListLevel | None) -> Length | None:
        """How far in this paragraph starts, adding up every reason there is.

        A numbered item is the one case where nothing has to be written: its
        indent comes from the numbering definition. That stops being true the
        moment anything else indents it too, because an explicit indent
        *replaces* the numbering's rather than adding to it — so once there is
        a quote or a definition in the way, the list's own step has to be
        counted back in by hand.
        """
        quoted = int(quote_indent(self._theme)) * self._quotes
        extra = int(self._extra_indent or 0)
        listed = int(self._lists[-1].indent) if self._lists else 0
        if item is not None and item.number is not None and not quoted and not extra:
            return None
        total = quoted + extra + listed
        return Emu(total) if total else None

    def _column_width(self) -> Length | None:
        """The width of the text column, which an image may not exceed.

        The last section rather than the first: front matter opens a section
        of its own, and the body an image sits in is the one after it.
        """
        return text_width(self._docx.sections[-1])

    def _skip_html(self, token: Token, line: int | None = None) -> None:
        """Warn once per line that raw HTML was dropped, and carry on.

        The PDF passes HTML through, because its pipeline is HTML. Word has
        nowhere to put it. Warning once per line rather than once per token
        keeps a paragraph with an opening and a closing tag in it from
        producing two warnings about one construct.
        """
        if HTML_COMMENT.match(token.content):
            return
        if line is None and token.map:
            line = token.map[0] + 1
        if line in self._warned_html:
            return
        self._warned_html.add(line)
        where = f" (line {line})" if line is not None else ""
        self._warn(f"raw HTML is not converted to Word; skipped it{where}")


# --- token helpers --------------------------------------------------------


def _closing(tokens: list[Token], index: int, end: int) -> int:
    """The index of the token that closes the container opened at ``index``."""
    depth = 0
    for offset in range(index, end):
        depth += tokens[offset].nesting
        if depth == 0:
            return offset
    return end - 1


def _href(token: Token) -> str:
    """Where a link points, read off the token that opened it."""
    value = token.attrGet("href")
    return value if isinstance(value, str) else ""


def _start(token: Token) -> int:
    """Where an ordered list is told to start counting."""
    value = token.attrGet("start")
    try:
        return int(str(value))
    except (TypeError, ValueError):
        return 1


def _footnote_number(token: Token) -> int:
    """Footnotes are numbered from zero in the token stream and one on paper."""
    return int(token.meta.get("id", 0)) + 1


def _table_rows(
    tokens: list[Token], start: int, end: int
) -> list[tuple[bool, list[tuple[Token | None, str | None]]]]:
    """Read a table's tokens into rows of cells, with the header row marked."""
    rows: list[tuple[bool, list[tuple[Token | None, str | None]]]] = []
    header = False
    cells: list[tuple[Token | None, str | None]] = []
    for index in range(start, end):
        token = tokens[index]
        if token.type == "thead_open":
            header = True
        elif token.type == "thead_close":
            header = False
        elif token.type == "tr_open":
            cells = []
        elif token.type == "tr_close":
            rows.append((header, cells))
        elif token.type in {"th_open", "td_open"}:
            content = tokens[index + 1] if index + 1 < end else None
            alignment = token.attrGet("style")
            cells.append(
                (
                    content
                    if content is not None and content.type == "inline"
                    else None,
                    alignment if isinstance(alignment, str) else None,
                )
            )
    return rows


def _why_unusable(exc: BaseException) -> str:
    """Why an image could not be embedded, in the words a person would use."""
    if isinstance(exc, OSError):
        return "it could not be read"
    if isinstance(exc, InvalidImageStreamError | UnexpectedEndOfFileError):
        return "the file is damaged, or is not the picture it claims to be"
    return "Word has no way to hold that image"


def _fit(picture: Any, available: Length | None) -> None:
    """Scale a picture down to the text column, keeping its proportions."""
    if available is None or picture.width <= available:
        return
    height = int(picture.height * int(available) / int(picture.width))
    picture.width = available
    picture.height = Emu(height)


#: Which character formatting each inline token turns on and off.
_OPENS = {"strong_open": "bold", "em_open": "italic", "s_open": "strike"}
_CLOSES = {"strong_close": "bold", "em_close": "italic", "s_close": "strike"}

#: The block tokens the walk acts on. Everything else — the closing halves,
#: and the tokens that only mark structure another handler already consumed —
#: is stepped over.
_BLOCKS = {
    "heading_open": _Builder._heading,
    "paragraph_open": _Builder._paragraph_block,
    "fence": _Builder._code,
    "code_block": _Builder._code,
    "blockquote_open": _Builder._quote,
    "bullet_list_open": _Builder._list,
    "ordered_list_open": _Builder._list,
    "list_item_open": _Builder._item,
    "hr": _Builder._rule,
    "dt_open": _Builder._term,
    "dd_open": _Builder._definition,
    "table_open": _Builder._table,
    "html_block": _Builder._html,
    "footnote_block_open": _Builder._footnote_block,
    "footnote_open": _Builder._footnote,
}


__all__ = ["render_docx"]
